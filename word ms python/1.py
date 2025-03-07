import docx

doc1 = docx.Document("1.docx")
doc2 = docx.Document("destination.docx")

# получаем из первого документа стили всех абзацев
styles = []
for paragraph in doc1.paragraphs:
    styles.append(paragraph.style)

# применяем стили ко всем абзацам второго документа
for i in range(len(doc2.paragraphs)):
    doc2.paragraphs[i].style = styles[i]

doc2.save("restored.docx")
